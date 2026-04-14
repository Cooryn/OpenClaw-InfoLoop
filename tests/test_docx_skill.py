from __future__ import annotations

from pathlib import Path

from docx import Document

from skills import docx_skill


def test_generate_report_creates_docx(tmp_path: Path) -> None:
    output = tmp_path / "report.docx"
    analyzed = [
        {
            "title": "Sample Title",
            "category": "Category A",
            "publication_date": "2026-04-14",
            "summary": "A short summary.",
            "url": "https://example.com/item",
            "content": "Full text content.",
        }
    ]

    report_path = docx_skill.generate_report(analyzed, str(output))
    assert Path(report_path).exists()
    assert Path(report_path).stat().st_size > 0

    loaded = Document(report_path)
    all_text = "\n".join(p.text for p in loaded.paragraphs)
    assert "Sample Title" in all_text
    assert "A short summary." in all_text


def test_add_text_watermark_empty_text_returns_false() -> None:
    document = Document()
    assert docx_skill.add_text_watermark(document, text="   ") is False

